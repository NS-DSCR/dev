import os
import logging
import io
import pandas as pd
from typing import List, Dict, Optional, Any
from abc import ABC, abstractmethod
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from rank_bm25 import BM25Okapi
from utils import get_embeddings
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class BaseKnowledgeBase(ABC):
    @abstractmethod
    def search(self, query: str, max_sections: int = 3) -> str:
        pass

class PersistentHybridKB(BaseKnowledgeBase):
    """
    Advanced Knowledge Base with:
    1. Persistent storage (ChromaDB)
    2. Hybrid Search (Vector Similarity + BM25 Keyword)
    3. Multi-format support
    """
    def __init__(self, agent_id: str, persist_directory: str = "./kb_storage", embedding_model: Optional[str] = None):
        self.agent_id = agent_id
        self.persist_dir = os.path.join(persist_directory, agent_id)
        self.embedding_model = embedding_model
        self.vector_store = None
        self.bm25 = None
        self.chunks: List[Document] = []
        
        # Initialize splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""]
        )
        
        # Try to load existing
        embeddings = get_embeddings(self.embedding_model)
        if os.path.exists(self.persist_dir) and embeddings:
            try:
                self.vector_store = Chroma(
                    persist_directory=self.persist_dir,
                    embedding_function=embeddings
                )
                # Rebuild BM25 from stored docs
                all_docs = self.vector_store.get()
                if all_docs['documents']:
                    self.chunks = [Document(page_content=d, metadata=m) for d, m in zip(all_docs['documents'], all_docs['metadatas'])]
                    tokenized_corpus = [doc.page_content.lower().split() for doc in self.chunks]
                    self.bm25 = BM25Okapi(tokenized_corpus)
                logger.info(f"💾 Loaded persistent KB for agent {agent_id}")
            except Exception as e:
                logger.error(f"Failed to load persistent KB: {e}")

    def ingest_content(self, content: str, source_name: str = "upload"):
        """Ingests raw text or markdown"""
        docs = [Document(page_content=content, metadata={"source": source_name})]
        self._process_and_store(docs)

    def ingest_binary(self, file_content: bytes, filename: str):
        """Processes PDF, Docx, or Excel bytes"""
        try:
            text = ""
            if filename.endswith('.pdf'):
                reader = PdfReader(io.BytesIO(file_content))
                text = "\n".join([page.extract_text() for page in reader.pages])
            elif filename.endswith('.docx'):
                doc = DocxDocument(io.BytesIO(file_content))
                text = "\n".join([para.text for para in doc.paragraphs])
            elif filename.endswith('.xlsx') or filename.endswith('.xls'):
                df = pd.read_excel(io.BytesIO(file_content))
                text = df.to_string()
            else:
                text = file_content.decode('utf-8', errors='ignore')

            if text:
                docs = [Document(page_content=text, metadata={"source": filename})]
                self._process_and_store(docs)
                return True
        except Exception as e:
            logger.error(f"Encoding/Parsing error for {filename}: {e}")
        return False

    def _process_and_store(self, docs: List[Document]):
        """Splits, embeds, and indexes documents"""
        split_docs = self.text_splitter.split_documents(docs)
        self.chunks.extend(split_docs)
        
        # 1. Update Chroma
        embeddings = get_embeddings(self.embedding_model)
        if not embeddings:
            logger.error("No embedding model configured")
            return

        if self.vector_store:
            self.vector_store.add_documents(split_docs)
        else:
            self.vector_store = Chroma.from_documents(
                documents=split_docs,
                embedding=embeddings,
                persist_directory=self.persist_dir
            )
        
        # 2. Update BM25
        tokenized_corpus = [doc.page_content.lower().split() for doc in self.chunks]
        self.bm25 = BM25Okapi(tokenized_corpus)
        logger.info(f"✅ Indexed {len(split_docs)} new chunks for agent {self.agent_id}")

    def search(self, query: str, max_sections: int = 3) -> str:
        if not self.chunks or not self.vector_store:
            return "Knowledge base is empty."

        # 1. Vector Search
        vector_results = self.vector_store.similarity_search(query, k=max_sections)
        
        # 2. BM25 Search
        tokenized_query = query.lower().split()
        bm25_scores = self.bm25.get_scores(tokenized_query)
        bm25_top_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:max_sections]
        bm25_results = [self.chunks[i] for i in bm25_top_indices if bm25_scores[i] > 0]

        # 3. Re-rank / Combine (Simplified: Unique union of both)
        combined_results = []
        seen_content = set()
        
        for doc in (vector_results + bm25_results):
            if doc.page_content not in seen_content:
                combined_results.append(doc)
                seen_content.add(doc.page_content)

        final_results = combined_results[:max_sections]
        if not final_results:
            return "No relevant information found in knowledge base."

        formatted_output = "\n\n".join([
            f"--- Source: {r.metadata.get('source', 'Unknown')} ---\n{r.page_content}\n" 
            for r in final_results
        ])
        
        return f"Found relevant info from various sources for '{query}':\n{formatted_output}"

# Compatibility wrapper for old interface
class MarkdownKnowledgeBase(BaseKnowledgeBase):
    def __init__(self, content: str = "", source_name: str = "default"):
        self.kb = PersistentHybridKB("legacy_" + source_name)
        if content: self.kb.ingest_content(content, source_name)
    def search(self, query: str, max_sections: int = 3) -> str:
        return self.kb.search(query, max_sections)

class VectorKnowledgeBase(BaseKnowledgeBase):
    def __init__(self, content: str = "", source_name: str = "default"):
        self.kb = PersistentHybridKB("vector_" + source_name)
        if content: self.kb.ingest_content(content, source_name)
    def search(self, query: str, max_sections: int = 3) -> str:
        return self.kb.search(query, max_sections)

class RemoteKnowledgeBase(BaseKnowledgeBase):
    def __init__(self, provider: str, url: str, api_key: str, source_name: str = "remote"):
        self.provider = provider
        self.url = url
        self.api_key = api_key
        self.source_name = source_name

    def search(self, query: str, max_sections: int = 3) -> str:
        import requests
        try:
            if self.provider == "rest_api":
                response = requests.post(self.url, headers={"Authorization": f"Bearer {self.api_key}"}, json={"query": query}, timeout=10)
                if response.status_code == 200:
                    results = response.json().get("results", [])
                    return "\n\n".join([r.get("text", "") for r in results[:max_sections]])
            return f"{self.provider.capitalize()} search for '{query}'"
        except Exception as e:
            return f"Remote Search Error: {e}"

# Global registry with persistence support
_KB_REGISTRY: Dict[str, BaseKnowledgeBase] = {}

def register_agent_kb(agent_id: str, content: str, kb_type: str = "markdown", **kwargs):
    """Registers a persistent hybrid KB by default now"""
    if kb_type == "remote":
        url = kwargs.get("kb_url")
        if url:
            _KB_REGISTRY[agent_id] = RemoteKnowledgeBase(kwargs.get("kb_provider", "rest_api"), url, kwargs.get("kb_api_key", ""), source_name=f"Agent-{agent_id}")
        return

    # For local KBs, we use the Hybrid Persistent version
    kb = PersistentHybridKB(agent_id, embedding_model=kwargs.get("kb_embedding_model"))
    if content:
        kb.ingest_content(content, "manual_input")
    
    _KB_REGISTRY[agent_id] = kb

def get_agent_kb(agent_id: str) -> Optional[BaseKnowledgeBase]:
    return _KB_REGISTRY.get(agent_id)
